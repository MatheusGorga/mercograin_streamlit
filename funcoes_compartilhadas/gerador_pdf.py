# Dentro de funcoes_compartilhadas/gerador_pdf.py

import docx
from fpdf import FPDF
from datetime import datetime

def criar_contrato_pdf(template_path, background_image_path, dados_contrato):
    """
    Gera um contrato em PDF preenchendo um template .docx e aplicando um fundo.
    
    :param template_path: Caminho para o arquivo .docx do template.
    :param background_image_path: Caminho para a imagem de fundo.
    :param dados_contrato: Dicionário com {placeholder: valor} para substituição.
    :return: Bytes do PDF gerado.
    """
    try:
        # --- 1. Preencher o template do Word ---
        document = docx.Document(template_path)

        for paragraph in document.paragraphs:
            for key, value in dados_contrato.items():
                placeholder = f"<<{key}>>"
                text_to_replace = str(value) if value is not None else ""
                
                if placeholder in paragraph.text:
                    inline = paragraph.runs
                    for i in range(len(inline)):
                        if placeholder in inline[i].text:
                            text = inline[i].text.replace(placeholder, text_to_replace)
                            inline[i].text = text
                            break

        texto_completo = [p.text for p in document.paragraphs]

        # --- 2. Criar o PDF com FPDF2 ---
        pdf = FPDF()
        pdf.add_page()
        
        pdf.image(background_image_path, x=0, y=0, w=00, h=100)
        
        pdf.set_y(20)
        pdf.set_x(20)
        
        pdf.set_font("Arial", size=11)

        for paragrafo in texto_completo:
            pdf.multi_cell(170, 5, txt=paragrafo, align='J')
            pdf.ln(3)

        # ✅ CORREÇÃO FINAL APLICADA AQUI ✅
        # Convertemos o 'bytearray' para o formato 'bytes' que o Streamlit espera.
        return bytes(pdf.output())

    except FileNotFoundError as e:
        raise FileNotFoundError(f"Arquivo de template ou imagem não encontrado: {e}")
    except Exception as e:
        raise Exception(f"Ocorreu um erro ao gerar o PDF: {e}")